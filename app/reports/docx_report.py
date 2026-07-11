from docx import Document

def generate_docx_report(report: dict, output_path: str, vendor_name: str = "Unknown Vendor"):
    doc = Document()
    doc.add_heading("Cloud Security Compliance Report", 0)
    doc.add_paragraph(f"Vendor: {vendor_name}")
    doc.add_paragraph(f"Total checks: {report['summary']['total']} | "
                       f"Passed: {report['summary']['passed']} | "
                       f"Failed: {report['summary']['failed']}")
    doc.add_heading("Detailed Results", level=1)
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Rule ID", "Resource", "Status", "Detail"
    for r in report["results"]:
        row = table.add_row().cells
        row[0].text, row[1].text, row[2].text, row[3].text = r["rule_id"], r["resource"], r["status"], r["detail"]
    doc.save(output_path)